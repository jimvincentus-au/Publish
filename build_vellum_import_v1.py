#!/usr/bin/env python3
"""
build_vellum_import_v1.py

Build Vellum-ready DOCX files (one per week) from canonical Markdown appendix artifacts.

Contract (v1):
- Input:  Publish/Output/Appendices/Week_XX/<week appendix>.md (tolerant discovery)
- Output: Publish/Output/Vellum/Appendices/Week_XX/Democracy_Clock_Year_One_Week_XX_Appendix.docx

Semantic mapping:
- Week title            -> Heading 1
- Category header       -> Heading 2
- Event line "N. ..."   -> List Number
- Summary paragraphs    -> Normal
"""

import argparse
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from typing import cast

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# --- Paths (match Publish pipeline conventions) ---
PUBLISH_ROOT = Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Publish")
PUBLISH_LOGS_DIR = PUBLISH_ROOT / "Logs"

DEFAULT_INPUT_ROOT = PUBLISH_ROOT / "Output" / "Scrivener"
DEFAULT_OUTPUT_ROOT = PUBLISH_ROOT / "Output" / "Vellum"
# For appendix discovery, prefer Substack markdown output for stable YAML front matter.
DEFAULT_SUBSTACK_ROOT = PUBLISH_ROOT / "Output" / "Substack"


# Preferred appendix JSON root (authoritative provenance: dates + sources)
DEFAULT_STEP3_WEEKS_ROOT = Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Step 3/Weeks")

LOG_NAME = "vellum_import"


def _resolve_input_root(override: Optional[str]) -> Path:
    if override:
        p = Path(override).expanduser()
        return p

    # Prefer Substack markdown output for appendices because it contains stable YAML
    # front matter (title/subtitle/week). Fall back to Scrivener output.
    candidates = [
        DEFAULT_SUBSTACK_ROOT,
        DEFAULT_INPUT_ROOT,
    ]

    for c in candidates:
        if c.exists() and c.is_dir():
            return c

    raise FileNotFoundError(
        f"Could not resolve appendix input root at expected location: {DEFAULT_INPUT_ROOT}. " 
        "This project assumes Scrivener output as the canonical source. " 
        "Use --input-root only if the project structure has changed."
    )


def _resolve_output_root(override: Optional[str]) -> Path:
    if override:
        return Path(override).expanduser()
    return DEFAULT_OUTPUT_ROOT


# --- Canonical Appendix Categories (MUST NOT DRIFT) ---
CANONICAL_CATEGORIES_ORDER: List[str] = [
    "Power and Authority",
    "Institutions and Governance",
    "Civil Rights and Dissent",
    "Economic Structure",
    "Information, Memory, and Manipulation",
]

# Legacy/alternate headings that must be mapped into the canonical names above.
# NOTE: Keys must match the cleaned heading text produced by _normalize_category_label.
CATEGORY_ALIASES: Dict[str, str] = {
    # canonical self-maps
    "Power and Authority": "Power and Authority",
    "Institutions and Governance": "Institutions and Governance",
    "Civil Rights and Dissent": "Civil Rights and Dissent",
    "Civil Rights & Dissent": "Civil Rights and Dissent",
    "Economic Structure": "Economic Structure",
    "Information, Memory, and Manipulation": "Information, Memory, and Manipulation",
    # common punctuation/connector variants that appear in some appendix JSON
    "Information, Memory and Manipulation": "Information, Memory, and Manipulation",
    "Information, Memory, and Manipulation": "Information, Memory, and Manipulation",
    "Information, Memory & Manipulation": "Information, Memory, and Manipulation",
    "Information, Memory and Manipulation ": "Information, Memory, and Manipulation",

    # older appendix naming (map into the canonical five)
    "Law and Justice": "Institutions and Governance",
    "Elections and Consent": "Civil Rights and Dissent",
    "Equality and Inclusion": "Civil Rights and Dissent",
    "Truth and Information": "Information, Memory, and Manipulation",
}


# --- Source display names (short tags -> formal/human names) ---
# Keep this mapping tight and deterministic; expand as new feeds are added.
SOURCE_DISPLAY_NAMES: Dict[str, str] = {
    # Core feeds
    "guardian": "The Guardian",
    "hcr": "Letters from an American (Heather Cox Richardson)",
    "popinfo": "Popular Information (Judd Legum)",

    # Democracy Clock internal / companion feeds
    "meidas": "Meidas Plus",
    "zeteo": "Zeteo",
    "noah": "Noahpinion (Noah Smith)",
    "outloud": "Democracy Outloud (Karen Zeigler)",
    "50501": "The 50501 Movement",

    # Government / official sources
    "orders": "White House / Executive Orders",
    "fr": "Federal Register",
    "federalregister": "Federal Register",
    "congress": "Congress.gov",
    "justsecurity": "Just Security",
    "scotus": "Supreme Court of the United States",
    "dhs": "U.S. Department of Homeland Security",
    "doj": "U.S. Department of Justice",
    "shadow": "Supreme Court Shadow Docket",
    "dod": "U.S. Department of Defense",

    # Media wires / outlets
    "nyt": "The New York Times",
    "wapo": "The Washington Post",
    "reuters": "Reuters",
    "ap": "Associated Press",
    "bbc": "BBC News",

    # Other
    "noaa": "NOAA",
}

def _format_source_names(keys: List[str]) -> List[str]:
    out: List[str] = []
    for k in keys:
        kk = str(k).strip()
        if not kk:
            continue
        out.append(SOURCE_DISPLAY_NAMES.get(kk, kk))
    # De-dup, preserve order
    seen = set()
    dedup: List[str] = []
    for s in out:
        if s in seen:
            continue
        seen.add(s)
        dedup.append(s)
    return dedup


def _format_iso_date(d: str) -> str:
    """Convert YYYY-MM-DD to 'D Mon YYYY' (e.g., 2025-01-22 -> 22 Jan 2025)."""
    s = str(d).strip()
    m = re.fullmatch(r"(20\d{2})-(\d{2})-(\d{2})", s)
    if not m:
        return s
    y, mo, da = int(m.group(1)), int(m.group(2)), int(m.group(3))
    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    if 1 <= mo <= 12:
        return f"{da} {months[mo-1]} {y}"
    return s


def _format_dates(dates: List[str]) -> List[str]:
    out = [_format_iso_date(x) for x in dates]
    # De-dup, preserve order
    seen = set()
    dedup: List[str] = []
    for s in out:
        if s in seen:
            continue
        seen.add(s)
        dedup.append(s)
    return dedup


def _parse_iso_date_for_sort(d: str) -> Optional[Tuple[int, int, int]]:
    """Parse YYYY-MM-DD into (y,m,d) for sorting. Returns None if not ISO."""
    s = str(d).strip()
    m = re.fullmatch(r"(20\d{2})-(\d{2})-(\d{2})", s)
    if not m:
        return None
    return int(m.group(1)), int(m.group(2)), int(m.group(3))


def _format_date_range(iso_dates: List[str]) -> Optional[str]:
    """Format an inclusive min–max date range from ISO dates (YYYY-MM-DD)."""
    parsed: List[Tuple[Tuple[int, int, int], str]] = []
    for d in iso_dates:
        p = _parse_iso_date_for_sort(d)
        if p is None:
            continue
        parsed.append((p, d))
    if not parsed:
        return None

    parsed.sort(key=lambda t: t[0])
    start_iso = parsed[0][1]
    end_iso = parsed[-1][1]

    if start_iso == end_iso:
        return _format_iso_date(start_iso)

    m1 = re.fullmatch(r"(20\d{2})-(\d{2})-(\d{2})", start_iso)
    m2 = re.fullmatch(r"(20\d{2})-(\d{2})-(\d{2})", end_iso)
    if not (m1 and m2):
        return f"{_format_iso_date(start_iso)}–{_format_iso_date(end_iso)}"

    y1, mo1, d1 = int(m1.group(1)), int(m1.group(2)), int(m1.group(3))
    y2, mo2, d2 = int(m2.group(1)), int(m2.group(2)), int(m2.group(3))
    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

    if y1 == y2 and mo1 == mo2 and 1 <= mo1 <= 12:
        return f"{d1}–{d2} {months[mo1-1]} {y1}"

    if y1 == y2 and 1 <= mo1 <= 12 and 1 <= mo2 <= 12:
        return f"{d1} {months[mo1-1]}–{d2} {months[mo2-1]} {y1}"

    return f"{_format_iso_date(start_iso)}–{_format_iso_date(end_iso)}"


def setup_logger(level: str) -> logging.Logger:
    """Create a simple file+console logger (mirrors build_wordpress_import_v1 style)."""
    PUBLISH_LOGS_DIR.mkdir(parents=True, exist_ok=True)
    log_path = PUBLISH_LOGS_DIR / "build_vellum_import_v1.log"

    logger = logging.getLogger(LOG_NAME)
    logger.setLevel(logging.DEBUG if level == "debug" else logging.INFO)

    if not logger.handlers:
        fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")

        fh = logging.FileHandler(log_path, encoding="utf-8")
        fh.setFormatter(fmt)
        logger.addHandler(fh)

        ch = logging.StreamHandler()
        ch.setFormatter(fmt)
        logger.addHandler(ch)

    return logger


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Build Vellum-ready DOCX appendix files from canonical Markdown artifacts"
    )
    p.add_argument("--week", type=int, required=True, help="Starting week number (N)")
    p.add_argument("--weeks", type=int, required=True, help="Number of weeks to include (K)")
    p.add_argument("--level", choices=["info", "debug"], default="info", help="Logging level (default: info)")
    p.add_argument(
        "--input-root",
        type=str,
        default=None,
        help="Override input root folder that contains week appendix markdown folders",
    )
    p.add_argument(
        "--output-root",
        type=str,
        default=None,
        help="Override output root folder for Vellum DOCX outputs",
    )
    return p.parse_args()


def _week_dir_names(week: int) -> List[str]:
    """Return common Week folder naming variants used across pipelines."""
    w2 = f"{week:02d}"
    return [
        f"Week {w2}",
        f"Week {week}",
        f"Week_{w2}",
        f"Week_{week}",
        f"Week{w2}",
        f"Week{week}",
        f"Week-{w2}",
        f"Week-{week}",
    ]


def _discover_week_md(input_root: Path, week: int) -> Path:
    """
    Discover the appendix markdown file for the week with tolerant patterns.
    Hard error if ambiguous or missing.
    """
    week_dir: Optional[Path] = None
    assert input_root is not None
    tried_dirs: List[Path] = []

    for name in _week_dir_names(week):
        cand = input_root / name
        tried_dirs.append(cand)
        if cand.exists() and cand.is_dir():
            week_dir = cand
            break

    if week_dir is None:
        # Fallback: any directory under input_root that contains the week number and starts with 'Week'
        # (e.g., 'Week_01', 'Week 01', 'Week01', etc.)
        fallback = []
        for p in input_root.glob("Week*"):
            if p.is_dir() and re.search(rf"\b0?{week}\b", p.name):
                fallback.append(p)
        if len(fallback) == 1:
            week_dir = fallback[0]
        elif len(fallback) > 1:
            names = ", ".join(x.name for x in sorted(fallback))
            raise RuntimeError(
                f"Ambiguous appendix input folder for Week {week:02d} under {input_root}. Found: {names}. "
                "Standardize Week folder naming or adjust discovery rules."
            )
        else:
            try:
                top = [p.name for p in sorted(input_root.iterdir()) if p.is_dir()][:30]
                logger = logging.getLogger(LOG_NAME)
                logger.debug(f"Top-level dirs under input_root: {top}")
            except Exception:
                pass
            tried = ", ".join(str(p) for p in tried_dirs)
            raise FileNotFoundError(
                f"Missing appendix input folder for Week {week:02d} under {input_root}. Tried: {tried}"
            )

    assert week_dir is not None
    # Common patterns observed / expected
    # Prefer Substack appendix md naming when present.
    candidates = [
        # Substack canonical appendix markdown
        week_dir / f"week{week:02d}_appendix_substack.md",
        week_dir / f"week{week}_appendix_substack.md",
        week_dir / f"week{week:02d}-appendix_substack.md",
        week_dir / f"week{week}-appendix_substack.md",

        # Scrivener/legacy appendix markdown
        week_dir / f"week{week:02d}_appendix.md",
        week_dir / f"week{week}_appendix.md",
        week_dir / f"week{week:02d}-appendix.md",
        week_dir / f"week{week}-appendix.md",
        week_dir / f"appendix_week{week:02d}.md",
        week_dir / f"appendix_week{week}.md",
    ]

    for c in candidates:
        if c.exists() and c.is_file():
            return c

    # Fallback: search for a single plausible appendix markdown file
    # Prefer substack appendix md
    patterns = [
        # Prefer substack appendix md
        f"*appendix*substack*week{week:02d}*.md",
        f"*appendix*substack*week{week}*.md",
        f"*week{week:02d}*appendix*substack*.md",
        f"*week{week}*appendix*substack*.md",
        f"*appendix*_substack*.md",
        f"*appendix*substack*.md",

        # Then legacy appendix files
        f"*appendix*week{week:02d}*.md",
        f"*appendix*week{week}*.md",
        f"*week{week:02d}*appendix*.md",
        f"*week{week}*appendix*.md",
        "*appendix*.md",
    ]

    found: List[Path] = []
    for pat in patterns:
        found.extend([p for p in week_dir.glob(pat) if p.is_file()])

    # De-dup
    found = sorted(list({p.resolve() for p in found}))

    if len(found) == 0:
        raise FileNotFoundError(
            f"Could not find appendix markdown for Week {week:02d} under {week_dir}. "
            f"Tried: {', '.join(str(x.name) for x in candidates)} and fallback globs."
        )

    # If multiple, require explicit disambiguation by tightening patterns upstream.
    if len(found) > 1:
        names = ", ".join(p.name for p in found)
        raise RuntimeError(
            f"Ambiguous appendix markdown for Week {week:02d} under {week_dir}. Found: {names}. "
            "Ensure only one appendix .md exists per week or standardize the filename."
        )

    return found[0]


def _discover_week_appendix_json(week: int, logger: logging.Logger) -> Optional[Path]:
    """
    Preferred: locate Step 7 output JSON for the week:
      /Step 3/Weeks/Week N/events_appendix_weekN.json
    Returns None if missing.
    """
    week_dir = DEFAULT_STEP3_WEEKS_ROOT / f"Week {week}"
    cand = week_dir / f"events_appendix_week{week}.json"
    if cand.exists() and cand.is_file():
        return cand

    cand2 = week_dir / f"events_appendix_week{week:02d}.json"
    if cand2.exists() and cand2.is_file():
        return cand2

    logger.debug(f"No appendix JSON found for Week {week:02d} under {week_dir}; will fall back to markdown.")
    return None


def _output_docx_path(output_root: Path, md_path: Path) -> Path:
    """
    Mirror the input filename, producing '<weekNN>_vellum.docx' directly under output_root.

    Example:
      input:  week04_appendix.md
      output: /Output/Vellum/week04_vellum.docx
    """
    output_root.mkdir(parents=True, exist_ok=True)

    stem = md_path.stem  # e.g., "week04_appendix"
    m = re.match(r"^(week\d{1,3})", stem, flags=re.IGNORECASE)
    base = m.group(1).lower() if m else stem.lower()

    filename = f"{base}_vellum.docx"
    return output_root / filename


def _normalize_category_label(line: str) -> str:
    """
    Normalize possible Markdown category header lines so they can be matched
    against CANONICAL_CATEGORIES.

    Examples handled:
      "## Power and Authority" -> "Power and Authority"
      "**Power and Authority**" -> "Power and Authority"
      "###   Law and Justice" -> "Law and Justice"
    """
    s = line.strip()

    # Strip leading markdown header markers
    s = re.sub(r"^\s*#{1,6}\s*", "", s).strip()

    # Strip surrounding emphasis markers (common in compiled markdown)
    # Remove repeated leading/trailing * or _ characters
    s = re.sub(r"^[\*_]+\s*", "", s)
    s = re.sub(r"\s*[\*_]+$", "", s)

    s = s.strip()

    # Strip common trailing punctuation that often appears in headings
    s = re.sub(r"[:\-–—\s]+$", "", s).strip()

    return s


def _canonicalize_category(line: str) -> Optional[str]:
    """
    Map a raw markdown line to a canonical category name (or None).
    Handles exact matches and prefix matches (e.g., 'Law and Justice — Courts').
    """
    s = _normalize_category_label(line)
    if not s:
        return None

    # Exact match
    if s in CATEGORY_ALIASES:
        return CATEGORY_ALIASES[s]

    # Prefix match (handles suffix descriptors)
    for k, v in CATEGORY_ALIASES.items():
        if s.startswith(k):
            return v

    return None


def _is_category_header(line: str) -> bool:
    return _canonicalize_category(line) is not None


# Appendix event lines are expected to be in the canonical log form:
#   "N. <summary> — <source(s)> — <domain> — <one-line relevance>"
# We therefore require at least one dash separator (em dash preferred) to avoid
# mistakenly treating numbered sublists (e.g., source lists) as events.
_EVENT_RE = re.compile(r"^\s*(\d+)\.\s+(.*)$")


def _is_event_line(text: str) -> bool:
    """Return True if the line looks like a top-level appendix event line."""
    s = text.strip()
    m = _EVENT_RE.match(s)
    if not m:
        return False

    rest = m.group(2).strip()
    # Require at least one separator that is characteristic of the event-log format.
    # Em dash is canonical; allow spaced hyphen as a fallback.
    if "—" in rest or " - " in rest:
        return True

    return False


def _looks_like_week_title(line: str) -> bool:
    # Expect: "# Week 1 Appendix: ..."
    s = line.strip()
    return s.startswith("#") and "Appendix" in s and "Week" in s


def _strip_md_h1(line: str) -> str:
    # "# Title" -> "Title"
    return line.lstrip("#").strip()


def _split_md_paragraphs(lines: List[str]) -> List[str]:
    """
    Convert markdown lines into paragraph blocks.
    Blank lines separate paragraphs.
    """
    paras: List[str] = []
    buf: List[str] = []
    for raw in lines:
        s = raw.rstrip()
        if not s.strip():
            if buf:
                paras.append(" ".join(x.strip() for x in buf).strip())
                buf = []
            continue
        buf.append(s)
    if buf:
        paras.append(" ".join(x.strip() for x in buf).strip())
    return [p for p in paras if p.strip()]


def _safe_list(x) -> List[str]:
    if x is None:
        return []
    if isinstance(x, list):
        return [str(i).strip() for i in x if str(i).strip()]
    if isinstance(x, str):
        s = x.strip()
        return [s] if s else []
    s = str(x).strip()
    return [s] if s else []


def _event_sources_dates(e: Dict) -> Tuple[List[str], List[str]]:
    """
    Return (sources, dates) using Step 7 enrichment fields when present.
    Accept multiple schema variants to be resilient.
    """
    prov = e.get("provenance") or {}
    sources = (
        _safe_list(prov.get("sources"))
        or _safe_list(e.get("source_names"))
        or _safe_list(e.get("sources"))
    )
    dates = (
        _safe_list(prov.get("dates"))
        or _safe_list(e.get("date"))
        or _safe_list(e.get("event_date"))
        or _safe_list(e.get("dates"))
    )
    return sources, dates


def _event_display_text(e: Dict) -> str:
    for k in ("text", "event", "description", "title", "summary"):
        v = e.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return str(e).strip()



def _format_appendix_event_text(e: Dict) -> str:
    """Produce the list-item text for Word/Vellum from an appendix event dict."""
    actor = e.get("actor")
    action = e.get("action")
    summary_line = e.get("summary_line")

    parts: List[str] = []

    # First sentence: actor + action, rendered as a sentence (must end with a period).
    first_sentence = ""
    if isinstance(actor, str) and actor.strip() and isinstance(action, str) and action.strip():
        first_sentence = f"{actor.strip()} {action.strip()}".strip()
    elif isinstance(action, str) and action.strip():
        first_sentence = action.strip()
    elif isinstance(actor, str) and actor.strip():
        # Unusual, but keep something useful
        first_sentence = actor.strip()

    if not first_sentence:
        # Last-resort fallback: use any reasonable textual field; never str(e)
        first_sentence = _event_display_text(e)

    # Ensure the first sentence ends with terminal punctuation.
    # Prefer a period, but do not add one after existing terminal punctuation.
    first_sentence = first_sentence.rstrip()
    if not first_sentence.endswith((".", "!", "?")):
        first_sentence = first_sentence.rstrip(" ")
        first_sentence = first_sentence.rstrip(".") + "."

    # Always start with the actor/action sentence.
    parts.append(first_sentence)

    # Optional second sentence
    if isinstance(summary_line, str) and summary_line.strip():
        parts.append(summary_line.strip())

    # Sources + dates (formatted)
    sources, dates = _event_sources_dates(e)
    sources_fmt = _format_source_names(sources)
    dates_fmt = _format_dates(dates)

    tail: List[str] = []
    if sources_fmt:
        tail.append("; ".join(sources_fmt))
    if dates_fmt:
        tail.append("; ".join(dates_fmt))

    if tail:
        parts.append("(" + " | ".join(tail) + ")")

    text = " ".join(p.strip() for p in parts if p and p.strip())
    return re.sub(r"\s+", " ", text).strip()




def _extract_week_summary_from_md(
    md_path: Path,
    logger: logging.Logger,
    subtitle_to_exclude: Optional[str] = None,
) -> Optional[str]:
    """Extract the first *summary* paragraph from the week appendix markdown.

    Substack appendix markdown usually contains:
      - H1 title
      - italic subtitle line
      - header image markdown
      - summary paragraph

    We want the first substantive paragraph that is NOT the subtitle and NOT the image line.
    """

    def _normalize_for_cmp(s: str) -> str:
        s2 = (s or "").strip()
        # Remove surrounding markdown italics markers
        if (s2.startswith("*") and s2.endswith("*")) or (s2.startswith("_") and s2.endswith("_")):
            s2 = s2[1:-1].strip()
        return re.sub(r"\s+", " ", s2).strip()

    subtitle_norm = _normalize_for_cmp(subtitle_to_exclude) if subtitle_to_exclude else None

    try:
        raw = md_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        logger.debug(f"Could not read markdown for summary extraction: {md_path} ({exc})")
        return None

    lines = raw.splitlines()
    idx = 0

    # Find the H1 title line
    while idx < len(lines) and not _looks_like_week_title(lines[idx]):
        idx += 1
    if idx >= len(lines):
        logger.debug(f"Summary extract: no recognizable week title/H1 found in {md_path}")
        return None

    idx += 1  # past title line

    summary_lines: List[str] = []
    while idx < len(lines):
        line = lines[idx].rstrip()

        # Stop at category header or explicit events header
        if _is_category_header(line) or re.fullmatch(
            r"Week\s+\d+\s+Events", line.strip(), flags=re.IGNORECASE
        ):
            break

        summary_lines.append(line)
        idx += 1

    paras = _split_md_paragraphs(summary_lines)
    logger.debug(
        f"Summary extract: candidate paragraphs={len(paras)} from {md_path.name}; subtitle_norm={subtitle_norm!r}"
    )

    for p in paras:
        p_norm = _normalize_for_cmp(p)
        if not p_norm:
            continue

        # Skip HTML comments
        if p_norm.startswith("<!--"):
            logger.debug(f"Summary extract: skipping html comment paragraph: {p_norm[:80]!r}")
            continue

        # Skip header image markdown
        if p_norm.startswith("!["):
            logger.debug(f"Summary extract: skipping image paragraph: {p_norm[:80]!r}")
            continue

        # Skip the subtitle paragraph if it matches
        if subtitle_norm and p_norm == subtitle_norm:
            logger.debug(f"Summary extract: skipping paragraph because it matches subtitle: {p_norm!r}")
            continue

        # Skip ultra-short noise
        if len(p_norm) < 25:
            logger.debug(f"Summary extract: skipping too-short paragraph: {p_norm!r}")
            continue

        logger.debug(f"Summary extract: selected summary paragraph: {p_norm[:120]!r}")
        return p_norm

    logger.debug(f"Summary extract: no suitable summary paragraph found in {md_path.name}")
    return None


# --- YAML front-matter extraction for short_title and subtitle ---
def _extract_front_matter(md_path: Path, logger: logging.Logger) -> Tuple[Optional[str], Optional[str]]:
    """Extract (short_title, subtitle) from YAML front matter in the appendix markdown.

    - short_title is derived from the text after 'Appendix:' in the front-matter title.
    - subtitle is taken from front-matter 'subtitle'.

    This parser is intentionally minimal (no external YAML dependency) and tolerant of:
      - Leading blank lines / BOM
      - Quoted values
      - Extra keys

    Returns (short_title, subtitle).
    """
    try:
        raw = md_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        logger.debug(f"Front matter: could not read markdown: {md_path} ({exc})")
        return None, None

    # Tolerate BOM and leading whitespace/blank lines
    raw0 = raw.lstrip("\ufeff")
    lines = raw0.splitlines()

    # Find first non-empty line
    i = 0
    while i < len(lines) and not lines[i].strip():
        i += 1

    if i >= len(lines) or lines[i].strip() != "---":
        first_nonempty = None
        for ln in lines:
            if ln.strip():
                first_nonempty = ln.strip()
                break
        logger.debug(
            f"Front matter: not present (no leading ---) for {md_path}; first_nonempty={first_nonempty!r}"
        )
        return None, None

    # Collect front matter lines until closing ---
    i += 1
    fm_lines: List[str] = []
    while i < len(lines):
        if lines[i].strip() == "---":
            break
        fm_lines.append(lines[i])
        i += 1

    if i >= len(lines):
        logger.debug(f"Front matter: unterminated (no closing ---) for {md_path}")
        return None, None

    title_val: Optional[str] = None
    subtitle_val: Optional[str] = None

    for line in fm_lines:
        if ":" not in line:
            continue
        k, v = line.split(":", 1)
        key = k.strip().lower()
        val = v.strip().strip('"').strip("'")
        if key == "title":
            title_val = val
        elif key == "subtitle":
            subtitle_val = val

    short_title: Optional[str] = None
    if title_val:
        if "Appendix:" in title_val:
            after = title_val.split("Appendix:", 1)[1].strip()
            if after:
                short_title = after
        else:
            # Fallback: if title doesn't contain Appendix:, still allow using text after 'Week N'
            short_title = None

    logger.debug(
        f"Front matter extracted for {md_path.name}: title={title_val!r}, short_title={short_title!r}, subtitle={subtitle_val!r}"
    )

    return short_title, subtitle_val

# --- Helper: Extract (short_title, subtitle) from appendix markdown (front matter or body) ---
def _extract_title_subtitle_from_md(md_path: Path, logger: logging.Logger) -> Tuple[Optional[str], Optional[str]]:
    """Extract (short_title, subtitle) from appendix markdown.

    Priority order:
      1) YAML front matter (title/subtitle)
      2) Body H1 line like '# Week N Appendix: <Short Title>'
      3) Body italic line immediately following title (for subtitle)

    Returns (short_title, subtitle).
    """
    # 1) YAML front matter
    short_title, subtitle = _extract_front_matter(md_path, logger)
    if short_title or subtitle:
        logger.debug(
            f"MD title/subtitle: using front matter for {md_path.name}: short_title={short_title!r}, subtitle={subtitle!r}"
        )
        return short_title, subtitle

    # 2) Fallback: parse from body
    try:
        raw = md_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        logger.debug(f"MD title/subtitle: could not read markdown: {md_path} ({exc})")
        return None, None

    lines = raw.splitlines()

    # Helpful debug context if extraction fails
    preview = " | ".join([ln.strip() for ln in lines[:20] if ln.strip()])
    logger.debug(f"MD title/subtitle: first non-empty lines preview for {md_path.name}: {preview}")

    i = 0
    # Skip leading blank lines
    while i < len(lines) and not lines[i].strip():
        i += 1

    # Skip any leading HTML comments (Scrivener/Substack markers)
    while i < len(lines) and lines[i].strip().startswith("<!--"):
        i += 1

    # Find first H1-ish line
    h1_idx: Optional[int] = None
    for j in range(i, min(len(lines), i + 80)):
        s = lines[j].strip()
        if not s:
            continue
        # H1 markdown
        if s.startswith("#"):
            h1_idx = j
            break
        # Occasionally title is plain text
        if re.match(r"^Week\s+\d+", s, flags=re.IGNORECASE):
            h1_idx = j
            break

    if h1_idx is None:
        logger.debug(f"MD title/subtitle: no H1/title line found for {md_path.name}")
        return None, None

    h1 = lines[h1_idx].strip()
    h1_clean = _strip_md_h1(h1)

    # Derive short title from patterns in the body title line
    body_short_title: Optional[str] = None
    if "Appendix:" in h1_clean:
        after = h1_clean.split("Appendix:", 1)[1].strip()
        if after:
            body_short_title = after
    elif ":" in h1_clean:
        after = h1_clean.split(":", 1)[1].strip()
        if after and not re.match(r"^\d{1,3}\s*$", after):
            body_short_title = after

    # Derive subtitle: first italic-only line after the title
    body_subtitle: Optional[str] = None
    k = h1_idx + 1
    # skip blanks, html comments, and header image markdown
    while k < len(lines):
        s = lines[k].strip()
        if not s:
            k += 1
            continue
        if s.startswith("<!--"):
            k += 1
            continue
        if s.startswith("!["):
            k += 1
            continue
        # Markdown italics line: *text* or _text_
        m = re.fullmatch(r"\*(.+)\*", s)
        if not m:
            m = re.fullmatch(r"_(.+)_", s)
        if m:
            candidate = m.group(1).strip()
            if candidate:
                body_subtitle = candidate
        break

    logger.debug(
        f"MD title/subtitle: body-derived for {md_path.name}: h1={h1_clean!r}, short_title={body_short_title!r}, subtitle={body_subtitle!r}"
    )

    return body_short_title, body_subtitle

def build_docx_from_md(md_path: Path, out_path: Path, logger: logging.Logger) -> None:
    raw = md_path.read_text(encoding="utf-8", errors="replace")
    lines = raw.splitlines()

    doc = Document()

    # --- New Title Logic ---
    # 1) Parse YAML front matter for short_title and subtitle (or fallback from body)
    short_title, subtitle = _extract_title_subtitle_from_md(md_path, logger)
    logger.debug(f"Week MD: extracted short_title={short_title!r}, subtitle={subtitle!r} from {md_path}")

    # 2) Discover week number from front matter or fallback to filename/discovery context
    week = None
    # Try to get week from front matter if present
    try:
        # Re-read YAML front matter for 'week:'
        if raw.lstrip().startswith("---"):
            fm_lines: List[str] = []
            src_lines = lines
            idx_fm = 1
            while idx_fm < len(src_lines):
                if src_lines[idx_fm].strip() == "---":
                    break
                fm_lines.append(src_lines[idx_fm])
                idx_fm += 1
            for line in fm_lines:
                if ":" not in line:
                    continue
                k, v = line.split(":", 1)
                key = k.strip().lower()
                val = v.strip().strip('"').strip("'")
                if key == "week":
                    try:
                        week = int(val)
                        break
                    except Exception:
                        pass
    except Exception:
        pass
    # Fallback: try to extract week number from filename if not found
    if week is None:
        m = re.search(r"week[_\- ]?(\d{1,3})", str(md_path.name), flags=re.IGNORECASE)
        if m:
            try:
                week = int(m.group(1))
            except Exception:
                week = None
    # Last resort: fallback to None (will display as "Week None")

    # 3) Compute date range using appendix JSON if available, else from event lines
    # Try to locate appendix JSON for provenance dates
    all_dates: List[str] = []
    try:
        from pathlib import Path as _Path
        appendix_json = _discover_week_appendix_json(week, logger) if week is not None else None
        if appendix_json:
            import json as _json
            data = _json.loads(appendix_json.read_text(encoding="utf-8", errors="replace"))
            raw_categories_for_dates = data.get("categories")
            if isinstance(raw_categories_for_dates, list):
                for c in raw_categories_for_dates:
                    if isinstance(c, dict):
                        evs = c.get("events") or []
                        if isinstance(evs, list):
                            for e in evs:
                                if isinstance(e, dict):
                                    _, ds = _event_sources_dates(e)
                                    all_dates.extend(ds)
            elif isinstance(raw_categories_for_dates, dict):
                for _, evs in raw_categories_for_dates.items():
                    if isinstance(evs, list):
                        for e in evs:
                            if isinstance(e, dict):
                                _, ds = _event_sources_dates(e)
                                all_dates.extend(ds)
    except Exception:
        all_dates = []
    # If no JSON provenance dates, fallback to dates parsed from Markdown event lines
    if not all_dates:
        # Parse dates from event lines in Markdown
        for line in lines:
            if _is_event_line(line):
                # Try to extract date(s) using heuristic: look for YYYY-MM-DD in line
                found = re.findall(r"\b20\d{2}-\d{2}-\d{2}\b", line)
                all_dates.extend(found)

    date_range_str = _format_date_range(all_dates) if all_dates else None

    # 4) Compose final title string (Heading 1)
    week_str = f"{week}" if week is not None else ""
    title_base = f"Week {week_str}"
    if date_range_str:
        title_base = f"{title_base} ({date_range_str})"
    if short_title:
        final_title = f"{title_base}: {short_title}"
    else:
        final_title = title_base
    logger.debug(f"Week MD: final chapter title={final_title!r}")
    doc.add_paragraph(final_title, style="Heading 1")

    # 5) If subtitle present, emit as indented italic blockquote
    if subtitle:
        p = doc.add_paragraph(subtitle, style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.10)
        for run in p.runs:
            run.italic = True
        logger.debug("Week MD: wrote subtitle as italic indented blockquote-style paragraph")

    # 6) Summary block: consume until we hit a category header (or explicit Week X Events header).
    # This must work for markdown that begins with YAML front matter (no H1) as well as legacy H1.
    idx = 0

    # Skip leading YAML front matter if present
    j = 0
    while j < len(lines) and not lines[j].strip():
        j += 1
    if j < len(lines) and lines[j].strip() == "---":
        j += 1
        while j < len(lines) and lines[j].strip() != "---":
            j += 1
        if j < len(lines) and lines[j].strip() == "---":
            j += 1
        idx = j
        logger.debug("Week MD: skipped YAML front matter for summary scanning")

    # Skip a legacy H1 title line if present
    if idx < len(lines) and _looks_like_week_title(lines[idx]):
        logger.debug(f"Week MD: skipping legacy title line for summary: {lines[idx].strip()!r}")
        idx += 1

    # Now buffer summary lines until the first category header / events header
    summary_lines: List[str] = []
    while idx < len(lines):
        line = lines[idx].rstrip()
        # Stop when we encounter a category heading or the explicit events header
        if _is_category_header(line) or re.fullmatch(r"Week\s+\d+\s+Events", line.strip(), flags=re.IGNORECASE):
            break
        summary_lines.append(line)
        idx += 1

    # Remove subtitle from summary lines if present, to avoid duplicate emission
    def _normalize_subtitle_cmp(s: str) -> str:
        s = s.strip()
        # Remove surrounding *...* or _..._ (Markdown italics)
        if (s.startswith("*") and s.endswith("*")) or (s.startswith("_") and s.endswith("_")):
            s = s[1:-1].strip()
        return s

    filtered_summary_lines = []
    subtitle_norm = _normalize_subtitle_cmp(subtitle) if subtitle else None
    for line in summary_lines:
        line_norm = _normalize_subtitle_cmp(line)
        if subtitle and line_norm == subtitle_norm:
            logger.debug(
                f"Summary: skipping subtitle line from summary emission: orig={line!r}, norm={line_norm!r}, subtitle_norm={subtitle_norm!r}"
            )
            continue
        filtered_summary_lines.append(line)

    # Write summary paragraphs (Normal)
    summary_paras = _split_md_paragraphs(filtered_summary_lines)
    if summary_paras:
        logger.debug(f"Summary paragraphs: {len(summary_paras)}")
        for p in summary_paras:
            doc.add_paragraph(p, style="Normal")
    else:
        logger.info("No summary paragraphs detected (continuing).")

    # 3) Advance past any explicit “Week X Events” header line(s), if present
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue
        # Stop at first category heading or first event line
        if _is_category_header(line) or _is_event_line(line):
            break
        # Skip only the explicit week events header
        if re.fullmatch(r"Week\s+\d+\s+Events", line, flags=re.IGNORECASE):
            logger.debug(f"Skipping events header line: {line}")
            idx += 1
            continue
        idx += 1

    # 4) Buffer all content by canonical category, then emit in canonical order
    buffered: Dict[str, List[Tuple[str, str]]] = {c: [] for c in CANONICAL_CATEGORIES_ORDER}
    categories_seen: List[str] = []
    current_category: Optional[str] = None
    in_category = False
    events_count = 0

    while idx < len(lines):
        raw_line = lines[idx]
        line = raw_line.strip()
        idx += 1

        if not line:
            continue

        # Category heading (canonicalize)
        cat = _canonicalize_category(line)
        if cat:
            prev = current_category
            current_category = cat
            in_category = True
            if cat not in categories_seen:
                categories_seen.append(cat)
            # Log only when the category changes, and include line context.
            if prev != cat:
                logger.debug(f"Category @L{idx}: {cat} (raw='{line}')")
            continue


        # Numbered event item (top-level appendix event line)
        if _is_event_line(line):
            m = _EVENT_RE.match(line)
            assert m is not None

            if not in_category or not current_category:
                # If events begin before any category heading, default to first canonical category
                current_category = CANONICAL_CATEGORIES_ORDER[0]
                in_category = True
                if current_category not in categories_seen:
                    categories_seen.append(current_category)
                logger.debug(f"Category (defaulted): {current_category}")

            # At this point, current_category is guaranteed to be a non-empty str
            assert current_category is not None

            # Remove native number; keep only the post-number text
            rest = m.group(2).strip()
            buffered[current_category].append(("list", rest))
            events_count += 1
            continue

        # Continuation / explanatory text line
        if in_category and current_category:
            buffered[current_category].append(("p", line))
        else:
            # Stray text outside category: attach to the first canonical category so nothing is lost
            buffered[CANONICAL_CATEGORIES_ORDER[0]].append(("p", line))

    # Emit categories in canonical order (always), even if empty
    for cat_name in CANONICAL_CATEGORIES_ORDER:
        doc.add_paragraph(cat_name, style="Heading 2")

        # Explicit numbering restart per category with a hanging indent.
        list_index = 1

        for kind, text in buffered[cat_name]:
            if kind == "list":
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.first_line_indent = Inches(-0.20)
                p.add_run(f"{list_index}. ")
                p.add_run(text)
                list_index += 1
            else:
                doc.add_paragraph(text, style="Normal")

    # 5) Verification logs
    logger.info(f"Input:  {md_path}")
    logger.info(f"Output: {out_path}")
    logger.info(f"Categories seen ({len(categories_seen)}): {', '.join(categories_seen) if categories_seen else '(none)'}")
    logger.info(f"Events emitted: {events_count}")
    if events_count > 200:
        logger.warning(
            f"Events emitted is unusually high for a single week ({events_count}). "
            "This often indicates the parser is still interpreting numbered sublists as events. "
            "Inspect the input markdown for non-event numbered lines and tighten _is_event_line() if needed."
        )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def build_docx_from_appendix_json(json_path: Path, out_path: Path, week: int, logger: logging.Logger) -> None:
    import json as _json

    data = _json.loads(json_path.read_text(encoding="utf-8", errors="replace"))

    # Heading 1 title
    title = None
    for k in ("title", "week_title", "heading"):
        v = data.get(k)
        if isinstance(v, str) and v.strip():
            title = v.strip()
            break
    # Remove fallback "Appendix" from title if not present
    if not title:
        title = f"Week {week:02d}"


    def _min_max_iso_dates(iso_dates: List[str]) -> Optional[Tuple[str, str]]:
        """Return (min_iso, max_iso) from a list of YYYY-MM-DD strings, ignoring non-ISO."""
        parsed: List[Tuple[Tuple[int, int, int], str]] = []
        for d in iso_dates:
            p = _parse_iso_date_for_sort(d)
            if p is None:
                continue
            parsed.append((p, str(d).strip()))
        if not parsed:
            return None
        parsed.sort(key=lambda t: t[0])
        return parsed[0][1], parsed[-1][1]

    # Compute deterministic min/max date range from provenance dates across all events
    all_dates: List[str] = []
    raw_categories_for_dates = data.get("categories")

    try:
        if isinstance(raw_categories_for_dates, list):
            for c in raw_categories_for_dates:
                if isinstance(c, dict):
                    evs = c.get("events") or []
                    if isinstance(evs, list):
                        for e in evs:
                            if isinstance(e, dict):
                                _, ds = _event_sources_dates(e)
                                all_dates.extend(ds)
        elif isinstance(raw_categories_for_dates, dict):
            for _, evs in raw_categories_for_dates.items():
                if isinstance(evs, list):
                    for e in evs:
                        if isinstance(e, dict):
                            _, ds = _event_sources_dates(e)
                            all_dates.extend(ds)
    except Exception:
        all_dates = []

    minmax = _min_max_iso_dates(all_dates)
    range_part = None
    if minmax:
        start_iso, end_iso = minmax
        range_part = f"{_format_iso_date(start_iso)} – {_format_iso_date(end_iso)}"

    # Derive short title from JSON fields or from common title patterns.
    # Priority:
    #  1) Explicit fields in JSON
    #  2) Text after 'Appendix:' in the title
    #  3) Text after the first ':' in the title
    short_title = None

    for k in ("short_title", "week_short_title", "week_title_short", "subtitle"):
        v = data.get(k)
        if isinstance(v, str) and v.strip():
            short_title = v.strip()
            break

    if not short_title and isinstance(title, str):
        t = title.strip()
        if "Appendix:" in t:
            after = t.split("Appendix:", 1)[1].strip()
            if after:
                short_title = after
        elif ":" in t:
            after = t.split(":", 1)[1].strip()
            if after:
                short_title = after

    # Compose final chapter title for Vellum/TOC
    base = f"Week {week}"
    if range_part:
        base = f"{base} ({range_part})"

    # Prefer short title from markdown front matter if available
    md_short_title = None
    md_subtitle = None
    md_path = None
    try:
        input_root = _resolve_input_root(None)
        md_path = _discover_week_md(input_root, week)
        logger.debug(f"Week {week:02d}: discovered markdown for title/subtitle: {md_path}")
        if md_path.exists():
            try:
                raw_preview = md_path.read_text(encoding="utf-8", errors="replace")
                preview_lines = [ln.rstrip() for ln in raw_preview.splitlines() if ln.strip()][:10]
                logger.debug(f"Week {week:02d}: md preview (first 10 non-empty lines): {preview_lines}")
                logger.debug(f"Week {week:02d}: md_path size={md_path.stat().st_size}")
            except Exception as _exc:
                logger.debug(f"Week {week:02d}: could not read md preview: {md_path} ({_exc})")
        else:
            logger.debug(f"Week {week:02d}: md_path does not exist: {md_path}")

        md_short_title, md_subtitle = _extract_title_subtitle_from_md(md_path, logger)
        logger.debug(f"Week {week:02d}: md-derived short_title={md_short_title!r}, subtitle={md_subtitle!r}")
    except Exception as exc:
        logger.debug(f"Week {week:02d}: could not derive short title/subtitle from markdown ({exc})")

    effective_short_title = md_short_title or short_title
    logger.debug(
        f"Week {week:02d}: title parts: base={base!r}, json_short_title={short_title!r}, md_short_title={md_short_title!r}, effective_short_title={effective_short_title!r}"
    )

    # Compose final chapter title (DROP the word 'Appendix')
    final_title = base
    if effective_short_title:
        final_title = f"{final_title}: {effective_short_title}"

    doc = Document()
    logger.debug(f"Week {week:02d}: final chapter title={final_title!r}")
    doc.add_paragraph(final_title, style="Heading 1")

    # Optional subtitle rendered as italic block paragraph
    if md_subtitle:
        p = doc.add_paragraph(md_subtitle, style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.10)
        for run in p.runs:
            run.italic = True
        logger.debug(f"Week {week:02d}: wrote subtitle as italic indented blockquote-style paragraph")

    # Optional summary paragraph (prefer JSON; fall back to markdown used for Substack)
    summary = data.get("summary")
    wrote_summary = False

    if isinstance(summary, str) and summary.strip():
        for ptxt in _split_md_paragraphs(summary.splitlines()):
            doc.add_paragraph(ptxt, style="Normal")
        wrote_summary = True
        logger.debug(f"Week {week:02d}: wrote summary paragraph(s) from JSON 'summary' field")

    if not wrote_summary:
        try:
            input_root = _resolve_input_root(None)
            md_path = _discover_week_md(input_root, week)
            logger.debug(
                f"Week {week:02d}: attempting markdown summary extraction from {md_path} (subtitle_to_exclude={md_subtitle!r})"
            )

            md_summary = _extract_week_summary_from_md(
                md_path,
                logger,
                subtitle_to_exclude=md_subtitle,
            )
            logger.debug(f"Week {week:02d}: markdown summary extractor returned: {md_summary!r}")

            if md_summary:
                doc.add_paragraph(md_summary, style="Normal")
                wrote_summary = True
                logger.debug(f"Week {week:02d}: wrote summary paragraph from markdown fallback")
            else:
                logger.debug(f"Week {week:02d}: no markdown-derived summary paragraph found")

        except Exception as exc:
            logger.debug(f"Week {week:02d}: Could not extract summary from markdown fallback ({exc})")

    raw_categories = data.get("categories")
    cat_to_events: Dict[str, List[Dict]] = {c: [] for c in CANONICAL_CATEGORIES_ORDER}

    def add_events(cat_name: str, events: List[Dict]) -> None:
        canon = _canonicalize_category(cat_name) or CATEGORY_ALIASES.get(_normalize_category_label(cat_name)) or cat_name
        canon = CATEGORY_ALIASES.get(canon, canon)
        if canon not in cat_to_events:
            logger.warning(
                f"Week {week:02d}: Unexpected category in appendix JSON: '{cat_name}' (normalized='{canon}'). "
                "Attaching events to first canonical category."
            )
            canon = CANONICAL_CATEGORIES_ORDER[0]
        cat_to_events[canon].extend([e for e in events if isinstance(e, dict)])

    if isinstance(raw_categories, list):
        for c in raw_categories:
            if not isinstance(c, dict):
                continue
            name = c.get("name") or c.get("category") or c.get("heading") or ""
            evs = c.get("events") or []
            if isinstance(name, str) and isinstance(evs, list):
                add_events(name, evs)
    elif isinstance(raw_categories, dict):
        for name, evs in raw_categories.items():
            if isinstance(name, str) and isinstance(evs, list):
                add_events(name, evs)

    # Emit with validation counters
    events_seen = 0
    events_with_sources = 0
    events_with_dates = 0

    # Track source-key coverage for deterministic reporting
    sources_seen_set = set()
    sources_unknown_set = set()

    for cat_name in CANONICAL_CATEGORIES_ORDER:
        doc.add_paragraph(cat_name, style="Heading 2")

        # Explicit numbering restart per category with a hanging indent.
        list_index = 1

        for e in cat_to_events[cat_name]:
            # provenance counters
            sources, dates = _event_sources_dates(e)
            # Track sources seen and unknowns
            for sk in sources:
                sk2 = str(sk).strip()
                if not sk2:
                    continue
                sources_seen_set.add(sk2)
                if sk2 not in SOURCE_DISPLAY_NAMES:
                    sources_unknown_set.add(sk2)
            events_seen += 1
            if sources:
                events_with_sources += 1
            if dates:
                events_with_dates += 1

            line = _format_appendix_event_text(e)

            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.20)
            p.add_run(f"{list_index}. ")
            p.add_run(line)
            list_index += 1

    # Week-level warnings (your “suspicious conditions”)
    if events_seen == 0:
        logger.warning(f"Week {week:02d}: Appendix JSON contains zero events: {json_path}")
    else:
        if events_with_dates == 0:
            logger.warning(f"Week {week:02d}: NO appendix events contain dates (JSON): {json_path}")
        elif events_with_dates < events_seen:
            logger.warning(f"Week {week:02d}: {events_seen - events_with_dates} appendix event(s) missing dates (JSON): {json_path}")

        if events_with_sources == 0:
            logger.warning(f"Week {week:02d}: NO appendix events contain sources (JSON): {json_path}")
        elif events_with_sources < events_seen:
            logger.warning(f"Week {week:02d}: {events_seen - events_with_sources} appendix event(s) missing sources (JSON): {json_path}")

    if sources_seen_set:
        seen_sorted = ", ".join(sorted(sources_seen_set))
        logger.info(f"Week {week:02d}: source keys seen: {seen_sorted}")
    if sources_unknown_set:
        unknown_sorted = ", ".join(sorted(sources_unknown_set))
        logger.warning(
            f"Week {week:02d}: UNKNOWN source key(s) with no formal mapping: {unknown_sorted}. "
            "Add them to SOURCE_DISPLAY_NAMES."
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    logger.info(
        f"Week {week:02d}: events={events_seen}, with_sources={events_with_sources}, with_dates={events_with_dates}"
    )
    logger.info(f"Appendix JSON input: {json_path}")


def main() -> None:
    args = parse_args()
    logger = setup_logger(args.level)

    logger.info("Starting Vellum DOCX build (appendix)")
    logger.debug(f"Args: week={args.week}, weeks={args.weeks}, level={args.level}")

    input_root = _resolve_input_root(args.input_root)
    output_root = _resolve_output_root(args.output_root)

    logger.info(f"Input root:  {input_root}")
    logger.info(f"Output root: {output_root}")
    try:
        logger.debug(f"Substack root exists={DEFAULT_SUBSTACK_ROOT.exists()} path={DEFAULT_SUBSTACK_ROOT}")
    except Exception:
        pass

    for week in range(args.week, args.week + args.weeks):
        logger.info(f"Processing Week {week:02d}")

        json_path = _discover_week_appendix_json(week, logger)

        if json_path:
            # Output name stays weekNN_vellum.docx even when input is JSON
            md_stub = Path(f"week{week:02d}_appendix.md")
            out_path = _output_docx_path(output_root, md_stub)

            logger.debug(f"Resolved json_path={json_path}")
            logger.debug(f"Resolved out_path={out_path}")

            build_docx_from_appendix_json(json_path, out_path, week, logger)
        else:
            md_path = _discover_week_md(input_root, week)
            out_path = _output_docx_path(output_root, md_path)

            logger.debug(f"Resolved md_path={md_path}")
            logger.debug(f"Resolved out_path={out_path}")

            build_docx_from_md(md_path, out_path, logger)

    logger.info("Done.")


if __name__ == "__main__":
    main()